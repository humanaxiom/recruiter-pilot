"""Unit tests for PDF/DOCX/RTF/TXT extraction (``src/pipeline/parsing/extract.py``).

Ported from hris ``packages/pipeline/src/pipeline/parsing/extract.py``
(``phase3-source-dossier.md`` §1). Pure functions over bytes — no live
services, no network. Real tiny PDF/DOCX fixtures are authored in-memory
with PyMuPDF (``fitz``) and ``python-docx`` rather than committed as
binaries; a hand-rolled fake ``fitz`` document is used where we need
precise control over block geometry (reading-order sort, image-block
filtering) that a real PDF layout can't reliably pin down.

The NUL-stripping tests guard a real production bug: Postgres ``text``/
``jsonb`` reject U+0000. An unstripped NUL crashes the DB write, leaves
the resume ``status='uploaded'``, and a reconcile cron re-queues it
forever, burning an LLM pass every time.

── Round 2 (merge-blocking re-audit) additions ─────────────────────────────
* Finding 1 (HIGH): the DOCX decompression-bomb guard must enforce a REAL
  decompression ceiling by streaming actual bytes, not by trusting the
  zip's self-declared (and forgeable) central-directory ``file_size``.
* Finding 2 (MAJOR): a malformed PDF that trips MuPDF's page-tree reader
  must surface as the typed ``UnsupportedMimeError``, never a bare
  ``RuntimeError``/internal MuPDF exception that escapes uncaught.
* Finding 5 (LOW): the 300/301-page cap boundary test already exists below
  (unchanged) — Finding 2's fix (wrapping the MuPDF page-tree access)
  subsumes the ``getattr(doc, "page_count", 0)`` fail-open.

── Round 3 (merge-blocking re-audit) additions ─────────────────────────────
* F5 (LOW): ``doc.needs_pass`` is read OUTSIDE any try block, between the
  ``fitz.open()`` try and the page-tree try. If that READ itself raises (a
  corrupt doc, not merely an encrypted one), the exception must surface as
  the typed ``UnsupportedMimeError`` — not escape ``extract_text`` uncaught
  — and the handle must still be closed.
"""

from __future__ import annotations

import struct
import zipfile
from collections.abc import Iterator
from io import BytesIO
from typing import Any

import fitz  # type: ignore[import-untyped]
import pytest
from docx import Document

from src.pipeline.parsing.extract import (
    MIME_DOCX,
    MIME_PDF,
    MIME_RTF,
    MIME_TXT,
    EncryptedPdfError,
    UnsupportedMimeError,
    extract_text,
)

# ── Finding 3 (HIGH): decompression bomb / unbounded parsing guard ─────────
#
# ``extract.py`` doesn't define this error yet — that's the whole point of
# these RED tests. Import it if it exists; otherwise fall back to a local
# placeholder so the rest of this module still collects and the ALREADY
# GREEN tests below stay green while only the new finding-3 tests go red.
# The coder is expected to name the real one ``InputTooLargeError`` (per the
# task spec) and raise it from ``extract_text`` for all three cases below.
try:
    from src.pipeline.parsing.extract import InputTooLargeError
except ImportError:  # pragma: no cover — RED until the coder adds it

    class InputTooLargeError(Exception):  # type: ignore[no-redef]
        """Placeholder so this module collects before the fix lands."""


# ── Real-fixture builders (no committed binaries) ───────────────────────────


def _make_simple_pdf_bytes(text: str) -> bytes:
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), text, fontsize=14)
    buf = BytesIO()
    doc.save(buf)
    doc.close()
    return buf.getvalue()


def _make_encrypted_pdf_bytes() -> bytes:
    doc = fitz.open()
    doc.new_page()
    buf = BytesIO()
    doc.save(
        buf,
        encryption=fitz.PDF_ENCRYPT_AES_256,
        owner_pw="owner-secret",
        user_pw="user-secret",
        permissions=int(fitz.PDF_PERM_PRINT),
    )
    doc.close()
    return buf.getvalue()


def _make_docx_bytes(paragraphs: list[str]) -> bytes:
    doc = Document()
    for para in paragraphs:
        doc.add_paragraph(para)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_pdf_with_pages(count: int) -> bytes:
    doc = fitz.open()
    for i in range(count):
        page = doc.new_page()
        page.insert_text((72, 72), f"page {i}", fontsize=10)
    buf = BytesIO()
    doc.save(buf)
    doc.close()
    return buf.getvalue()


def _make_bomb_docx_bytes(uncompressed_size: int) -> bytes:
    """A real zip whose declared (uncompressed) ``word/document.xml`` size is
    ``uncompressed_size`` bytes, but whose ON-DISK size is tiny — the content
    is a single repeated character, so DEFLATE compresses it to almost
    nothing while the zip's central directory still truthfully declares the
    full decompressed size. That's exactly the shape of a zip-bomb-ish DOCX:
    small to upload, huge to fully inflate.
    """
    payload = "A" * uncompressed_size
    buf = BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("[Content_Types].xml", "<Types/>")
        zf.writestr("word/document.xml", payload)
    return buf.getvalue()


def _patch_zip_member_declared_size(
    blob: bytes, member_name: str, declared_size: int
) -> bytes:
    """Rewrite BOTH the local file header's and the central directory
    record's uncompressed-size ("file_size") field for ``member_name`` to
    ``declared_size``, leaving the compressed bytes and CRC32 untouched. A
    real unzip of the member still yields its full, TRUE decompressed
    content — only the size the zip self-reports is a lie. This is exactly
    what fools a central-directory-trusting guard
    (``sum(info.file_size for info in zf.infolist())``).
    """
    data = bytearray(blob)
    name_bytes = member_name.encode("utf-8")

    with zipfile.ZipFile(BytesIO(bytes(data))) as zf:
        info = zf.getinfo(member_name)
    # Local file header: signature(4) + 18 bytes of fixed fields ->
    # uncompressed size sits at offset +22.
    struct.pack_into("<I", data, info.header_offset + 22, declared_size)

    # Central directory record: signature(4) + fixed fields -> uncompressed
    # size ("file_size") sits at +24, filename length at +28, filename at +46.
    sig = b"PK\x01\x02"
    idx = 0
    patched_cd = False
    while True:
        idx = data.find(sig, idx)
        if idx == -1:
            break
        name_len = struct.unpack_from("<H", data, idx + 28)[0]
        if bytes(data[idx + 46 : idx + 46 + name_len]) == name_bytes:
            struct.pack_into("<I", data, idx + 24, declared_size)
            patched_cd = True
            break
        idx += 4
    assert patched_cd, f"central directory record for {member_name!r} not found"
    return bytes(data)


def _make_bomb_docx_bytes_with_lying_header(
    actual_uncompressed_size: int, declared_size: int
) -> bytes:
    """A real zip whose ``word/document.xml`` member ACTUALLY decompresses to
    ``actual_uncompressed_size`` bytes, but whose local + central-directory
    headers are rewritten to declare only ``declared_size`` bytes. The CRC32
    is untouched (it matches the real payload), so a genuine unzip /
    ``Document()`` call still succeeds and yields the full, true content —
    the exact "lying header" decompression bomb security proved (198 MB peak
    allocation despite a declared 50 bytes).
    """
    payload = "0" * actual_uncompressed_size  # highly compressible: tiny on disk
    buf = BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("[Content_Types].xml", "<Types/>")
        zf.writestr("word/document.xml", payload)
    blob = _patch_zip_member_declared_size(
        buf.getvalue(), "word/document.xml", declared_size
    )
    with zipfile.ZipFile(BytesIO(blob)) as zf:
        forged = zf.getinfo("word/document.xml").file_size
        assert forged == declared_size
        assert forged < actual_uncompressed_size
    return blob


# ── Fake fitz document (precise control over blocks/geometry) ───────────────


class _FakePage:
    def __init__(self, text: str, blocks: list[tuple[Any, ...]]) -> None:
        self._text = text
        self._blocks = blocks

    def get_text(self, mode: str) -> Any:
        if mode == "text":
            return self._text
        if mode == "blocks":
            return self._blocks
        raise AssertionError(f"unexpected get_text mode: {mode!r}")


class _FakeDoc:
    def __init__(self, pages: list[_FakePage], *, needs_pass: bool = False) -> None:
        self._pages = pages
        self.needs_pass = needs_pass
        self.closed = False

    def __iter__(self) -> Iterator[_FakePage]:
        return iter(self._pages)

    def close(self) -> None:
        self.closed = True


def _patch_fitz_open(monkeypatch: pytest.MonkeyPatch, fake_doc: _FakeDoc) -> None:
    monkeypatch.setattr(
        "src.pipeline.parsing.extract.fitz.open", lambda **_kw: fake_doc
    )


# ── Dispatch / happy path ────────────────────────────────────────────────────


def test_extract_text_dispatches_pdf() -> None:
    blob = _make_simple_pdf_bytes("Hello PDF World unique-marker-123")
    result = extract_text(blob, MIME_PDF)
    assert result.method == "pymupdf"
    assert "unique-marker-123" in result.full_text


def test_extract_text_dispatches_docx() -> None:
    blob = _make_docx_bytes(["First paragraph line.", "Second paragraph line."])
    result = extract_text(blob, MIME_DOCX)
    assert result.method == "docx"
    assert result.pages[0].page_no == 0
    assert "First paragraph line." in result.full_text
    assert "Second paragraph line." in result.full_text


def test_extract_text_dispatches_txt() -> None:
    blob = ("plain text content " * 15).encode("utf-8")
    result = extract_text(blob, MIME_TXT)
    assert result.method == "text"
    assert result.pages[0].page_no == 0


def test_extract_text_dispatches_rtf() -> None:
    rtf = (
        rb"{\rtf1\ansi\deff0{\fonttbl{\f0 Arial;}}\f0 Hello RTF World "
        rb"and some extra padding text so the extraction has plenty of "
        rb"content to work with in this deterministic unit test case.}"
    )
    result = extract_text(rtf, MIME_RTF)
    assert result.method == "rtf"
    assert "Hello RTF World" in result.full_text


def test_extract_text_raises_on_unsupported_mime() -> None:
    with pytest.raises(UnsupportedMimeError):
        extract_text(b"whatever", "application/x-nonsense")


# ── Encrypted PDF ─────────────────────────────────────────────────────────


def test_encrypted_pdf_raises_encrypted_pdf_error() -> None:
    blob = _make_encrypted_pdf_bytes()
    with pytest.raises(EncryptedPdfError):
        extract_text(blob, MIME_PDF)


# ── NUL stripping (Postgres text/jsonb reject U+0000) ────────────────────────


def test_nul_bytes_stripped_from_txt_full_text() -> None:
    blob = ("clean text " * 30 + "\x00more text after a nul\x00").encode("utf-8")
    result = extract_text(blob, MIME_TXT)
    assert "\x00" not in result.full_text
    assert "\x00" not in result.pages[0].text


def test_nul_bytes_stripped_from_pdf_page_text_and_blocks(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    blocks = [(0.0, 0.0, 10.0, 10.0, "clean\x00block\x00text", 0, 0)]
    page = _FakePage(text="page text with\x00nul chars " * 20, blocks=blocks)
    _patch_fitz_open(monkeypatch, _FakeDoc([page]))

    result = extract_text(b"fake-pdf-bytes", MIME_PDF)

    assert "\x00" not in result.pages[0].text
    assert all("\x00" not in b.text for b in result.pages[0].blocks)
    assert "\x00" not in result.full_text


# ── `_decode` ladder: utf-8-sig -> utf-8 -> latin-1, never raises ───────────


def test_decode_ladder_never_raises_on_undecodable_bytes() -> None:
    blob = b"\xff\xfe\x80\x81\x82Not valid UTF-8 but must not raise."
    result = extract_text(blob, MIME_TXT)  # would raise UnicodeDecodeError if unguarded
    assert isinstance(result.full_text, str)


def test_decode_ladder_prefers_utf8_sig_bom() -> None:
    text = "café résumé " * 20
    blob = b"\xef\xbb\xbf" + text.encode("utf-8")
    result = extract_text(blob, MIME_TXT)
    assert result.full_text.startswith("café")


# ── PDF blocks: reading-order sort + image-block filter ─────────────────────


def test_pdf_blocks_sorted_by_y_then_x_and_image_blocks_filtered(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    blocks = [
        (50.0, 100.0, 150.0, 120.0, "second (lower)", 0, 0),
        (5.0, 10.0, 100.0, 30.0, "first (top-left)", 1, 0),
        (0.0, 0.0, 10.0, 10.0, "an image block", 2, 1),  # b[6] != 0 -> filtered
    ]
    page = _FakePage(text="first (top-left)\nsecond (lower)", blocks=blocks)
    _patch_fitz_open(monkeypatch, _FakeDoc([page]))

    result = extract_text(b"fake-pdf-bytes", MIME_PDF)

    texts = [b.text for b in result.pages[0].blocks]
    assert texts == ["first (top-left)", "second (lower)"]
    assert len(result.pages[0].blocks) == 2


# ── Low text density warning (< 200 total extracted chars) ──────────────────


def test_low_text_density_warning_appended_for_short_pdf(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    page = _FakePage(text="short", blocks=[])
    _patch_fitz_open(monkeypatch, _FakeDoc([page]))

    result = extract_text(b"fake-pdf-bytes", MIME_PDF)

    assert "low_text_density_consider_ocr" in result.warnings


def test_low_text_density_warning_for_short_txt() -> None:
    result = extract_text(b"short text", MIME_TXT)
    assert "low_text_density_consider_ocr" in result.warnings


def test_no_low_text_density_warning_for_long_txt() -> None:
    long_text = ("word " * 60).encode("utf-8")  # well over 200 chars
    result = extract_text(long_text, MIME_TXT)
    assert "low_text_density_consider_ocr" not in result.warnings


# ── Finding 3a: byte-cap — must reject BEFORE touching fitz/docx at all ─────

_BYTE_CAP = 10 * 1024 * 1024  # 10 MB — matches hris's store_uploaded_blob cap


def test_extract_text_raises_before_any_parse_when_blob_exceeds_byte_cap(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    def _boom(*_a: object, **_kw: object) -> Any:
        raise AssertionError("fitz/docx must not be touched before the byte cap check")

    monkeypatch.setattr("src.pipeline.parsing.extract.fitz.open", _boom)
    monkeypatch.setattr("src.pipeline.parsing.extract.Document", _boom)

    oversized = b"x" * (_BYTE_CAP + 1)

    with pytest.raises(InputTooLargeError):
        extract_text(oversized, MIME_PDF)


@pytest.mark.parametrize("mime", [MIME_PDF, MIME_DOCX, MIME_TXT, MIME_RTF])
def test_extract_text_raises_on_oversized_blob_for_every_mime(mime: str) -> None:
    oversized = b"x" * (_BYTE_CAP + 1)
    with pytest.raises(InputTooLargeError):
        extract_text(oversized, mime)


def test_extract_text_accepts_blob_at_exactly_the_byte_cap() -> None:
    """The cap is a MAX, not an off-by-one trap: exactly ``_BYTE_CAP`` bytes
    of plain text must still parse."""
    blob = b"a" * _BYTE_CAP
    result = extract_text(blob, MIME_TXT)
    assert isinstance(result.full_text, str)


# ── Finding 3b: DOCX decompression-bomb — declared size cap ─────────────────


def test_extract_text_raises_on_decompression_bomb_docx() -> None:
    # 100 MB of a single repeated byte: DEFLATE crushes this to a few KB on
    # disk (well under the 10 MB byte cap) but the zip's own central
    # directory truthfully declares ~100 MB of decompressed content — a
    # bounded-input check must catch that BEFORE lxml fully inflates it.
    blob = _make_bomb_docx_bytes(100 * 1024 * 1024)
    assert len(blob) < _BYTE_CAP  # sanity: bomb passes the byte cap, must be
    # caught by a SEPARATE decompressed-size check

    with pytest.raises(InputTooLargeError):
        extract_text(blob, MIME_DOCX)


# ── Finding 1 (HIGH, round 2): the guard must not trust a LYING header ──────
#
# Security proved 198 MB peak allocation despite a forged central-directory
# ``file_size`` of 50 bytes. ``_docx_decompressed_size`` today just sums the
# declared ``ZipInfo.file_size`` from the central directory — an attacker who
# controls the zip controls that number independently of the real
# decompressed payload. The guard must enforce a REAL decompression ceiling
# by streaming each member and aborting once cumulative ACTUAL decompressed
# bytes exceed the cap, not by trusting the header.


def test_extract_text_raises_on_docx_bomb_with_forged_small_declared_header() -> None:
    """``word/document.xml`` genuinely decompresses to 60 MB (> the 50 MB
    cap), but its local + central-directory headers LIE and declare only 50
    bytes. A CD-trusting guard reads the lie, passes, and hands the bomb to
    ``Document()`` — which then genuinely inflates the full 60 MB. This must
    still raise ``InputTooLargeError``."""
    blob = _make_bomb_docx_bytes_with_lying_header(
        actual_uncompressed_size=60_000_000, declared_size=50
    )
    with pytest.raises(InputTooLargeError):
        extract_text(blob, MIME_DOCX)


# ── Finding 3c: PDF page-count cap (300) ─────────────────────────────────────
#
# Finding 5 (LOW, round 2): these same two tests also pin the page-count
# cap's fail-open — ``getattr(doc, "page_count", 0)`` only intercepts
# ``AttributeError``, but a corrupt page tree makes the ``page_count``
# property RAISE ``RuntimeError`` instead, crashing the guard line itself.
# Finding 2's fix (wrapping the MuPDF page-tree access) subsumes this; no
# separate test is added here beyond the existing 300/301 boundary below.


def test_extract_text_raises_when_pdf_exceeds_page_cap() -> None:
    blob = _make_pdf_with_pages(301)
    with pytest.raises(InputTooLargeError):
        extract_text(blob, MIME_PDF)


def test_extract_text_accepts_pdf_at_exactly_the_page_cap() -> None:
    blob = _make_pdf_with_pages(300)
    result = extract_text(blob, MIME_PDF)
    assert len(result.pages) == 300


# ── Finding 2 (MAJOR, round 2): malformed PDF must raise a TYPED error ──────
#
# MuPDF-backed properties (``doc.page_count``) and page iteration execute C
# code that can raise an internal, UNTYPED exception on a malformed PDF —
# a bare ``RuntimeError`` (e.g. "code=7: Invalid number of pages") from the
# ``page_count`` property, or (confirmed against the real PyMuPDF used here)
# ``pymupdf.mupdf.FzErrorFormat`` ("cycle in page tree") from the per-page
# iteration loop. NEITHER is one of the typed errors (``UnsupportedMimeError``
# / ``EncryptedPdfError`` / ``InputTooLargeError``) that ``parse_resume``
# catches. Left unguarded, that exception escapes ``extract_text`` ->
# ``parse_resume`` uncaught, stranding the resume row with a NULL
# ``failure_reason`` and causing an arq retry storm on the same bytes
# forever. The fix wraps the MuPDF access broadly (``except Exception`` ->
# ``UnsupportedMimeError``), not a special case for one error string — these
# three DIFFERENT malformed-PDF shapes, tripping TWO different MuPDF code
# paths (``page_count`` access and per-page iteration), all must hit the
# same fix.

_MALFORMED_PDF_HUGE_COUNT_EMPTY_KIDS = (
    b"%PDF-1.4\n"
    b"1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n"
    b"2 0 obj\n<< /Type /Pages /Kids [] /Count 999999 >>\nendobj\n"
    b"trailer\n<< /Size 3 /Root 1 0 R >>\n"
    b"%%EOF"
)

# A Pages node whose own Kids array points back at itself. ``page_count``
# access succeeds, but the per-page iteration loop (`for i, page in
# enumerate(doc): ...` — wrapped only in a `finally`, never an `except`)
# detects the cycle and raises `pymupdf.mupdf.FzErrorFormat("cycle in page
# tree")`, a DIFFERENT untyped exception than Finding 2's other fixture, from
# a DIFFERENT code path (confirmed against the real PyMuPDF dependency
# pinned in requirements.txt).
_MALFORMED_PDF_CIRCULAR_PAGE_TREE = (
    b"%PDF-1.4\n"
    b"1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n"
    b"2 0 obj\n<< /Type /Pages /Kids [2 0 R] /Count 1 >>\nendobj\n"
    b"trailer\n<< /Size 3 /Root 1 0 R >>\n"
    b"%%EOF"
)

_MALFORMED_PDF_GARBAGE_AFTER_HEADER = b"%PDF-1.7\n" + (
    b"\x00\xff\xee not xref not obj just garbage bytes " * 5
)


@pytest.mark.parametrize(
    "bad_pdf",
    [
        _MALFORMED_PDF_HUGE_COUNT_EMPTY_KIDS,
        _MALFORMED_PDF_CIRCULAR_PAGE_TREE,
        _MALFORMED_PDF_GARBAGE_AFTER_HEADER,
    ],
    ids=["huge_count_empty_kids", "circular_page_tree", "garbage_after_header"],
)
def test_extract_text_raises_typed_error_not_bare_runtimeerror_on_malformed_pdf(
    bad_pdf: bytes,
) -> None:
    with pytest.raises(UnsupportedMimeError):
        extract_text(bad_pdf, MIME_PDF)


# ── F5 (LOW, round 3) — `doc.needs_pass` read failure must not escape ──────
# uncaught / leak the handle ────────────────────────────────────────────────
#
# `_extract_pdf` reads `doc.needs_pass` OUTSIDE any try block, between the
# `fitz.open()` try and the page-tree try. If that read itself raises on a
# corrupt doc (not merely an ENCRYPTED one — a genuinely malformed encryption
# dictionary), the exception must surface as the typed `UnsupportedMimeError`
# (never a bare/untyped exception escaping `extract_text` uncaught) and the
# handle must still be closed. The existing `EncryptedPdfError` path (above)
# must stay green — this only hardens the case where the CHECK ITSELF throws.


class _FakeDocNeedsPassRaises:
    """A doc whose `needs_pass` READ itself raises — simulating a corrupt PDF
    where even MuPDF's encryption-dictionary check throws, distinct from a
    genuinely password-protected doc (`EncryptedPdfError`, tested above)."""

    def __init__(self) -> None:
        self.closed = False

    @property
    def needs_pass(self) -> bool:
        raise RuntimeError("mupdf: corrupt encryption dictionary")

    def close(self) -> None:
        self.closed = True

    def __iter__(self) -> Iterator[Any]:
        return iter(())


def test_needs_pass_read_failure_raises_unsupported_mime_and_closes_handle(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    fake_doc = _FakeDocNeedsPassRaises()
    monkeypatch.setattr(
        "src.pipeline.parsing.extract.fitz.open", lambda **_kw: fake_doc
    )

    with pytest.raises(UnsupportedMimeError):
        extract_text(b"fake-pdf-bytes", MIME_PDF)

    assert fake_doc.closed is True, (
        "the fitz handle must be closed even when the needs_pass READ itself "
        "raises, not just when the page-tree loop raises"
    )
