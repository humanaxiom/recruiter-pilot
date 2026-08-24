"""The synthetic fixture generator must produce a corpus that is safe to publish
AND useful enough to rank.

Both halves matter and they pull against each other. A generator that emits
lorem is trivially safe and certifies a broken ranker as healthy — the exact trap
`model_probe_live` documents, where a synthetic prompt of the same length passed
while the real document failed 3/3. A generator that emits realistic contact
details is useful and reintroduces the disclosure this whole exercise exists to
end. So the tests pin both edges.
"""

from __future__ import annotations

import json
from pathlib import Path

import pytest

from src.synth_fixtures import (
    SYNTHETIC_MARKER,
    generate,
    load_vocabulary,
    person_for_jd,
    resume_lines,
    skills_for_jd,
    synthetic_marker_present,
    write_pdf,
)

_SKILL_DATA = Path(__file__).resolve().parents[2] / "src" / "pipeline" / "skill_data"

_JD_TEXT = (
    "The successful candidate will use python and sql to maintain reporting, "
    "with excel for budget tracking and project management across teams."
)


@pytest.fixture(scope="module")
def vocabulary() -> list[str]:
    return load_vocabulary(_SKILL_DATA)


# ── safe to publish ──────────────────────────────────────────────────────


def test_contact_details_use_reserved_ranges_only(vocabulary: list[str]) -> None:
    """RFC 2606 reserves `.invalid` as permanently unresolvable and 555-01xx is
    reserved for fiction. Any other domain or exchange risks generating a real
    person's contact details by coincidence — which is the entire failure this
    generator exists to prevent."""
    for jd in ("a.docx", "b.docx", "c.docx", "d.docx", "e.docx"):
        person = person_for_jd(jd, _JD_TEXT, vocabulary)
        assert person.email.endswith("@example.invalid")
        assert person.phone.startswith("555-01")


def test_every_generated_pdf_is_mechanically_identifiable(
    tmp_path: Path, vocabulary: list[str]
) -> None:
    """Filename conventions get renamed; metadata travels with the bytes. This
    is what lets an operator prove a fixture directory holds no real résumé
    rather than asserting it."""
    person = person_for_jd("x.docx", _JD_TEXT, vocabulary)
    out = tmp_path / "x_resume.pdf"
    write_pdf(resume_lines(person), out)
    assert synthetic_marker_present(out)


def test_the_marker_is_visible_in_the_document_body_too(
    vocabulary: list[str],
) -> None:
    """A human opening the PDF must be able to tell. Metadata is invisible in
    every viewer people actually use."""
    person = person_for_jd("y.docx", _JD_TEXT, vocabulary)
    assert any(SYNTHETIC_MARKER in line for line in resume_lines(person))


# ── useful enough to rank ────────────────────────────────────────────────


def test_skills_are_drawn_from_the_jd_not_invented(vocabulary: list[str]) -> None:
    """The load-bearing property. A résumé built from an arbitrary skill list
    ranks identically against every posting, so a corpus like that cannot tell a
    working ranker from a broken one."""
    skills = skills_for_jd(_JD_TEXT, vocabulary, coverage=1.0)
    assert skills, "JD mentions vocabulary terms; none were picked up"
    lowered = _JD_TEXT.lower()
    assert all(s.lower() in lowered for s in skills)


def test_skills_match_on_word_boundaries_not_substrings(
    vocabulary: list[str],
) -> None:
    """Found by reading the generated corpus, not by a test.

    The first version used ``skill in lowered`` and produced a Student Services
    Advisor whose skills were "go, scala, ssis" — ``ssis`` is inside *assist*,
    ``go`` inside *goals*, ``scala`` inside *escalate*. Short vocabulary entries
    turn a substring test into a false-positive generator, and a corpus of
    plausible-looking nonsense would have been measured against for weeks before
    anyone actually opened a résumé.
    """
    prose = (
        "The incumbent will assist with escalated enquiries, support ongoing "
        "goals, and coordinate assessment activities across the department."
    )
    picked = {s.lower() for s in skills_for_jd(prose, vocabulary, coverage=1.0)}
    for bogus in ("ssis", "go", "scala", "r", "c"):
        assert bogus not in picked, (
            f"{bogus!r} was matched inside a longer word — substring matching "
            "regressed"
        )


def test_a_genuinely_mentioned_skill_is_still_matched(vocabulary: list[str]) -> None:
    """The load-bearing other direction. Tightening to word boundaries must not
    become 'match nothing', which would pass the test above trivially."""
    picked = {s.lower() for s in skills_for_jd(_JD_TEXT, vocabulary, coverage=1.0)}
    assert "python" in picked
    assert "sql" in picked


def test_generated_text_avoids_glyphs_the_base_font_cannot_render(
    vocabulary: list[str],
) -> None:
    """PyMuPDF's built-in ``helv`` has no em-dash, en-dash or bullet, so those
    came out of the extractor as a fallback character — "2023·present". The
    parsed text is what the LLM sees, so its fidelity is the product's, not a
    cosmetic concern."""
    person = person_for_jd("glyphs.docx", _JD_TEXT, vocabulary)
    body = "\n".join(resume_lines(person))
    for glyph in ("—", "–", "•"):
        assert glyph not in body


def test_coverage_actually_reduces_the_skill_count(vocabulary: list[str]) -> None:
    """Coverage tiers are what make the pool rankable. If every tier returned
    the same list, every candidate would match everything and ordering could not
    be demonstrated."""
    full = skills_for_jd(_JD_TEXT, vocabulary, coverage=1.0)
    if len(full) < 3:
        pytest.skip("fixture JD text is too small to show a coverage spread")
    assert len(skills_for_jd(_JD_TEXT, vocabulary, coverage=0.3)) < len(full)


def test_a_jd_with_no_recognised_skills_still_yields_a_usable_resume(
    vocabulary: list[str],
) -> None:
    """The parser must still get a well-formed document — an empty résumé would
    fail extraction for reasons that have nothing to do with the code under
    test, and would look like a parser bug."""
    person = person_for_jd("empty.docx", "Nothing recognisable here.", vocabulary)
    lines = resume_lines(person)
    assert lines[0] == person.name
    assert any("SKILLS" in line for line in lines)


# ── deterministic ────────────────────────────────────────────────────────


def test_the_same_jd_always_produces_the_same_person(vocabulary: list[str]) -> None:
    """Seeded from the JD filename via sha256, not `hash()` — which is salted
    per process, so a `hash()`-seeded corpus would differ between runs and
    between machines. A fixture corpus that changes under you is not a fixture
    corpus."""
    first = person_for_jd("stable.docx", _JD_TEXT, vocabulary)
    second = person_for_jd("stable.docx", _JD_TEXT, vocabulary)
    assert (first.name, first.email, first.phone, first.skills) == (
        second.name,
        second.email,
        second.phone,
        second.skills,
    )


def test_different_jds_produce_different_people(vocabulary: list[str]) -> None:
    names = {
        person_for_jd(f"jd_{i}.docx", _JD_TEXT, vocabulary).name for i in range(12)
    }
    assert len(names) > 1, "seeding collapsed — every JD produced the same person"


# ── end to end ───────────────────────────────────────────────────────────


def test_generate_writes_one_resume_per_jd_with_a_manifest(tmp_path: Path) -> None:
    import docx

    jd_dir = tmp_path / "JDs"
    jd_dir.mkdir()
    for name in ("alpha", "beta", "gamma"):
        doc = docx.Document()
        doc.add_paragraph(_JD_TEXT)
        doc.save(str(jd_dir / f"{name}.docx"))

    out = tmp_path / "resumes"
    manifest = generate(jd_dir, out, _SKILL_DATA)

    pdfs = sorted(out.glob("*_resume.pdf"))
    assert len(pdfs) == 3
    assert len(manifest) == 3
    assert all(synthetic_marker_present(p) for p in pdfs)

    written = json.loads((out / "manifest.json").read_text(encoding="utf-8"))
    assert {row["source_jd"] for row in written} == {
        "alpha.docx",
        "beta.docx",
        "gamma.docx",
    }
    assert all(row["synthetic"] == "true" for row in written)


def test_generated_names_match_the_smoke_suite_glob(tmp_path: Path) -> None:
    """`tests/smoke/conftest.py` globs `resumes/*_resume.pdf` and asserts it
    finds three. A generator that named files anything else would leave the
    smoke suite failing for a reason unrelated to the product."""
    import docx

    jd_dir = tmp_path / "JDs"
    jd_dir.mkdir()
    for name in ("one", "two", "three"):
        doc = docx.Document()
        doc.add_paragraph(_JD_TEXT)
        doc.save(str(jd_dir / f"{name}.docx"))

    out = tmp_path / "resumes"
    generate(jd_dir, out, _SKILL_DATA)
    assert len(sorted(out.glob("*_resume.pdf"))) == 3


def test_temporary_word_lock_files_are_not_treated_as_jds(tmp_path: Path) -> None:
    """Word leaves `~$name.docx` lock files beside real documents. Feeding one
    to python-docx raises, and a fixture generator that dies because someone had
    a JD open is a bad tool."""
    import docx

    jd_dir = tmp_path / "JDs"
    jd_dir.mkdir()
    doc = docx.Document()
    doc.add_paragraph(_JD_TEXT)
    doc.save(str(jd_dir / "real.docx"))
    (jd_dir / "~$real.docx").write_bytes(b"not a real docx")

    manifest = generate(jd_dir, tmp_path / "out", _SKILL_DATA)
    assert len(manifest) == 1
