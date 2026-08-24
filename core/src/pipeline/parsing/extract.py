"""PDF / DOCX / RTF / TXT text extraction.

Primary path uses PyMuPDF (`fitz`) for PDFs and `python-docx` for
DOCX. OCR fallback (Tesseract) is documented and stubbed but not
implemented this phase.

Ported near-verbatim from hris
``packages/pipeline/src/pipeline/parsing/extract.py`` (see
``phase3-source-dossier.md`` §1). Only the import paths and the
structlog -> stdlib logging conversion changed.
"""

from __future__ import annotations

import logging
import zipfile
from dataclasses import dataclass, field
from io import BytesIO
from typing import Literal

import fitz  # type: ignore[import-untyped]
from docx import Document

log = logging.getLogger(__name__)

MIME_PDF = "application/pdf"
MIME_DOCX = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
MIME_RTF = "application/rtf"
MIME_TXT = "text/plain"

ExtractionMethod = Literal["pymupdf", "docx", "rtf", "text", "ocr"]

# ── Bounded input (this module is the TRUST BOUNDARY) ───────────────────────
#
# extract_text() is the first code to touch attacker-supplied bytes, so it has
# to be safe STANDALONE — it cannot lean on an upload-side byte cap that does
# not exist yet (Phase 6). The worker runs max_jobs=4, so an unbounded parse is
# four concurrent bombs away from taking the box down.
#
# Each cap is a MAXIMUM: a blob of exactly _MAX_BLOB_BYTES, or a PDF of exactly
# _MAX_PDF_PAGES, is ACCEPTED.
_MAX_BLOB_BYTES = 10 * 1024 * 1024  # matches hris's store_uploaded_blob cap
_MAX_PDF_PAGES = 300
# A 10 MB DOCX at a 1000:1 compression ratio inflates to ~10 GB inside lxml.
# The zip's central directory declares the decompressed size of every member up
# front, so we can reject the bomb WITHOUT inflating a single byte of it.
_MAX_DOCX_DECOMPRESSED_BYTES = 50 * 1024 * 1024


class UnsupportedMimeError(ValueError):
    """Mime type isn't one of the allowed inputs."""


class EncryptedPdfError(RuntimeError):
    """PDF requires a password we don't have."""


class InputTooLargeError(ValueError):
    """Input exceeds a hard bound: raw bytes, PDF pages, or — for DOCX — the
    size it would occupy once decompressed (the zip-bomb case)."""


@dataclass(frozen=True)
class TextBlock:
    text: str
    bbox: tuple[float, float, float, float]
    page_no: int


@dataclass(frozen=True)
class PageText:
    page_no: int
    text: str
    blocks: tuple[TextBlock, ...] = ()


@dataclass(frozen=True)
class ExtractedText:
    pages: tuple[PageText, ...]
    method: ExtractionMethod
    warnings: tuple[str, ...] = field(default_factory=tuple)

    @property
    def full_text(self) -> str:
        return "\n".join(p.text for p in self.pages)

    @property
    def total_chars(self) -> int:
        return sum(len(p.text) for p in self.pages)


# ---------------- public ----------------


def extract_text(blob: bytes, mime: str) -> ExtractedText:
    """Sync extractor — both backends are CPU-bound and small.

    Async callers should run this in a thread pool
    (`asyncio.to_thread(extract_text, blob, mime)`).

    Raises ``InputTooLargeError`` for a blob over ``_MAX_BLOB_BYTES``, checked
    BEFORE any parser is handed the bytes.
    """
    if len(blob) > _MAX_BLOB_BYTES:
        raise InputTooLargeError(
            f"blob is {len(blob)} bytes; the cap is {_MAX_BLOB_BYTES}"
        )
    if mime == MIME_PDF:
        result = _extract_pdf(blob)
    elif mime == MIME_DOCX:
        result = _extract_docx(blob)
    elif mime == MIME_RTF:
        result = _extract_rtf(blob)
    elif mime == MIME_TXT:
        result = _extract_txt(blob)
    else:
        raise UnsupportedMimeError(f"unsupported mime: {mime!r}")
    return _sanitize(result)


def _sanitize(result: ExtractedText) -> ExtractedText:
    """Strip NUL (U+0000) from extracted text. Postgres ``text``/``jsonb`` cannot
    store U+0000 (it raises ``UntranslatableCharacterError: unsupported Unicode
    escape sequence`` at write time), and some PDFs extract embedded NULs. Left
    unstripped, the résumé parse crashes at the DB write, the row stays
    ``uploaded``, and the reconcile cron re-queues it forever (burning an LLM
    pass each time). Rebuild the frozen dataclasses only when a NUL is
    actually present, so the common clean path is a no-op."""

    def _has_nul(p: PageText) -> bool:
        return "\x00" in p.text or any("\x00" in b.text for b in p.blocks)

    if not any(_has_nul(p) for p in result.pages):
        return result
    pages = tuple(
        PageText(
            page_no=p.page_no,
            text=p.text.replace("\x00", ""),
            blocks=tuple(
                TextBlock(
                    text=b.text.replace("\x00", ""), bbox=b.bbox, page_no=b.page_no
                )
                for b in p.blocks
            ),
        )
        for p in result.pages
    )
    return ExtractedText(pages=pages, method=result.method, warnings=result.warnings)


def _decode(blob: bytes) -> str:
    """Tolerant byte decode: UTF-8 BOM -> UTF-8 -> latin-1 (never fails)."""
    for codec in ("utf-8-sig", "utf-8"):
        try:
            return blob.decode(codec)
        except UnicodeDecodeError:
            continue
    return blob.decode("latin-1")


# ---------------- PDF ----------------


def _extract_pdf(blob: bytes) -> ExtractedText:
    """PyMuPDF extraction.

    Blocks are sorted by (y, x) so two-column layouts collapse into a
    sensible reading order. Image-only blocks (the b[6] == 1 case in
    fitz's block tuple) are skipped — they'd only contribute to the
    OCR path which is deferred.
    """
    try:
        doc = fitz.open(stream=blob, filetype="pdf")
    except Exception as exc:
        raise UnsupportedMimeError(f"pdf parse failed: {exc}") from exc

    # ``doc.needs_pass`` runs MuPDF C code that reads the encryption dictionary;
    # on a CORRUPT (not merely password-protected) doc that READ itself can
    # raise an untyped exception. Wrap it so a genuine password prompt still
    # surfaces as EncryptedPdfError, but a failed check surfaces as the typed
    # UnsupportedMimeError — never escaping extract_text uncaught — and the
    # handle is closed either way.
    try:
        needs_pass = doc.needs_pass
    except Exception as exc:
        doc.close()
        raise UnsupportedMimeError(f"pdf encryption check failed: {exc}") from exc
    if needs_pass:
        doc.close()
        raise EncryptedPdfError("pdf is password-protected")

    pages: list[PageText] = []
    warnings: list[str] = []
    # ``doc.page_count`` AND the page-iteration loop both run MuPDF C code that,
    # on a malformed page tree, raises an UNTYPED exception NOT in our caught
    # family — a bare ``RuntimeError`` (``code=7: Invalid number of pages``) from
    # the page-count read, or ``pymupdf.mupdf.FzErrorFormat`` (``cycle in page
    # tree``) from the loop. Any of them means "this PDF is unparseable", so both
    # are wrapped broadly into the typed ``UnsupportedMimeError`` rather than
    # allowed to escape ``extract_text`` -> ``parse_resume`` uncaught (a stranded
    # row + arq retry storm). ``InputTooLargeError`` (the page cap) and the
    # already-handled ``EncryptedPdfError`` stay distinct — the page cap is
    # re-raised, never swallowed. NOTE: ``getattr`` here only suppresses a MISSING
    # ``page_count`` attribute (real fitz always exposes it); a page_count
    # property that RAISES propagates through ``getattr`` and is caught by the
    # broad ``except`` below — so this is NOT a fail-open on a corrupt tree.
    try:
        page_count = int(getattr(doc, "page_count", 0) or 0)
        if page_count > _MAX_PDF_PAGES:
            raise InputTooLargeError(
                f"pdf has {page_count} pages; the cap is {_MAX_PDF_PAGES}"
            )
        for i, page in enumerate(doc):
            text = page.get_text("text") or ""
            # blocks = list of (x0, y0, x1, y1, text, block_no, block_type)
            raw_blocks = page.get_text("blocks") or []
            text_blocks = [
                TextBlock(
                    text=b[4],
                    bbox=(b[0], b[1], b[2], b[3]),
                    page_no=i,
                )
                for b in sorted(raw_blocks, key=lambda b: (b[1], b[0]))
                if len(b) > 6 and b[6] == 0  # text blocks only
            ]
            pages.append(PageText(page_no=i, text=text, blocks=tuple(text_blocks)))
    except InputTooLargeError:
        raise
    except Exception as exc:
        raise UnsupportedMimeError(f"pdf page tree unreadable: {exc}") from exc
    finally:
        doc.close()

    total = sum(len(p.text) for p in pages)
    if total < 200 and pages:
        warnings.append("low_text_density_consider_ocr")
        # TODO(phase-5b): rasterise + tesseract here. Today the worker
        # surfaces this as a soft warning and the LLM does its best
        # with whatever text we did get.

    return ExtractedText(
        pages=tuple(pages),
        method="pymupdf",
        warnings=tuple(warnings),
    )


# ---------------- DOCX ----------------


# Bounded chunk read for the streaming decompression guard — big enough that a
# legitimate DOCX inflates in a handful of iterations, small enough that a bomb
# never materialises whole in memory.
_DOCX_STREAM_CHUNK_BYTES = 1024 * 1024


def _enforce_docx_decompression_cap(blob: bytes) -> None:
    """Reject a DOCX that would inflate past ``_MAX_DOCX_DECOMPRESSED_BYTES``.

    Do NOT trust the zip's self-declared ``ZipInfo.file_size`` — the central
    directory is attacker-controlled and can LIE (a forged 50-byte declaration
    over an 80 MB member sails past a CD-summing guard, then ``Document()``
    inflates it anyway; security measured 198 MB peak). Instead enforce a REAL
    ceiling: open each member with ``zf.open()`` and read it in bounded chunks,
    summing ACTUAL decompressed bytes, and abort the instant the running total
    exceeds the cap — before ``Document()`` ever sees the bytes and without
    building any member whole in memory. ``zipfile.read()`` (unbounded) is never
    called on an untrusted member. A huge NUMBER of small members is handled by
    the same running sum for free. XXE is separately mitigated (python-docx sets
    ``resolve_entities=False``); this is the decompression-bomb half.
    """
    total = 0
    try:
        with zipfile.ZipFile(BytesIO(blob)) as zf:
            for info in zf.infolist():
                # Neutralise the self-declared uncompressed size FIRST. zipfile's
                # reader stops a member at ``ZipInfo.file_size`` bytes, so a
                # forged-small ``file_size`` (the lying-header bomb) would make
                # ``.read()`` truncate at the fake size and CRC-fail instead of
                # revealing the real payload. Overriding it with a large sentinel
                # makes ZipExtFile stop only at the true end of the DEFLATE
                # stream — but we abort far earlier, the instant the running total
                # of ACTUAL decompressed bytes crosses the cap, so a real bomb
                # never inflates whole in memory and never reaches the CRC check.
                info.file_size = 1 << 40
                with zf.open(info) as member:
                    while True:
                        chunk = member.read(_DOCX_STREAM_CHUNK_BYTES)
                        if not chunk:
                            break
                        total += len(chunk)
                        if total > _MAX_DOCX_DECOMPRESSED_BYTES:
                            raise InputTooLargeError(
                                "docx decompresses past the cap of "
                                f"{_MAX_DOCX_DECOMPRESSED_BYTES} bytes"
                            )
    except zipfile.BadZipFile as exc:
        raise UnsupportedMimeError(f"docx parse failed: {exc}") from exc


def _extract_docx(blob: bytes) -> ExtractedText:
    """python-docx — flat paragraph stream. DOCX has no real pages, so
    everything goes into page_no=0 and the UI surfaces 'Document' as
    the page label.
    """
    # Stream-verify the REAL decompressed size before Document() inflates it —
    # python-docx has no size guard of its own.
    _enforce_docx_decompression_cap(blob)

    try:
        doc = Document(BytesIO(blob))
    except Exception as exc:
        raise UnsupportedMimeError(f"docx parse failed: {exc}") from exc

    text = "\n".join(p.text for p in doc.paragraphs if p.text)
    return ExtractedText(
        pages=(PageText(page_no=0, text=text),),
        method="docx",
        warnings=(),
    )


# ---------------- RTF ----------------


def _extract_rtf(blob: bytes) -> ExtractedText:
    """striprtf — pure-Python RTF -> plain text. Single logical page like
    DOCX. RTF is 7-bit ASCII with escapes, so a tolerant decode is safe.

    striprtf is imported lazily so a process that never extracts RTF
    doesn't hard-require the optional dependency at import time."""
    from striprtf.striprtf import rtf_to_text

    try:
        text = rtf_to_text(_decode(blob)).strip()  # type: ignore[no-untyped-call]
    except Exception as exc:
        raise UnsupportedMimeError(f"rtf parse failed: {exc}") from exc
    warnings = ("low_text_density_consider_ocr",) if len(text) < 200 else ()
    return ExtractedText(
        pages=(PageText(page_no=0, text=text),), method="rtf", warnings=warnings
    )


# ---------------- plain text ----------------


def _extract_txt(blob: bytes) -> ExtractedText:
    """Plain text — tolerant decode, single page. No structure to extract."""
    text = _decode(blob).strip()
    warnings = ("low_text_density_consider_ocr",) if len(text) < 200 else ()
    return ExtractedText(
        pages=(PageText(page_no=0, text=text),), method="text", warnings=warnings
    )


__all__ = [
    "MIME_DOCX",
    "MIME_PDF",
    "MIME_RTF",
    "MIME_TXT",
    "EncryptedPdfError",
    "ExtractedText",
    "ExtractionMethod",
    "InputTooLargeError",
    "PageText",
    "TextBlock",
    "UnsupportedMimeError",
    "extract_text",
]
